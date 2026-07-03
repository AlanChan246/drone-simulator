#!/usr/bin/env python3
"""Generate the ADRC 2026 competition-day student handout (DOCX)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "docs" / "比賽當日參賽說明.docx"

FONT_NAME = "Microsoft JhengHei"
FONT_FALLBACK = "PingFang TC"


def set_run_font(run, size_pt, bold=False, color=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.insert(0, rFonts)


def add_paragraph(doc, text, size=10.5, bold=False, align=None, space_after=4, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 16, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6 if level > 1 else 0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run(text)
    set_run_font(run, sizes.get(level, 11), bold=True)
    return p


def add_bullet(doc, text, size=10.5, bold_parts=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.4)

    if not bold_parts:
        run = p.add_run(text)
        set_run_font(run, size)
        return p

    remaining = text
    for part, is_bold in bold_parts:
        idx = remaining.find(part)
        if idx == -1:
            continue
        if idx > 0:
            run = p.add_run(remaining[:idx])
            set_run_font(run, size)
        run = p.add_run(part)
        set_run_font(run, size, bold=is_bold)
        remaining = remaining[idx + len(part) :]
    if remaining:
        run = p.add_run(remaining)
        set_run_font(run, size)
    return p


def add_table(doc, headers, rows, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, 10, bold=True)

    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(cell_text)
            set_run_font(run, 10)

    if col_widths_cm:
        for row in table.rows:
            for i, width in enumerate(col_widths_cm):
                row.cells[i].width = Cm(width)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def shade_paragraph(paragraph, fill="E8F4FD"):
    pPr = paragraph._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)


def add_highlight_box(doc, lines, fill="E8F4FD"):
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2 if i < len(lines) - 1 else 4)
        p.paragraph_format.space_before = Pt(4 if i == 0 else 0)
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.line_spacing = 1.15
        shade_paragraph(p, fill)
        if isinstance(line, tuple):
            text, bold = line
            run = p.add_run(text)
            set_run_font(run, 10.5, bold=bold)
        else:
            run = p.add_run(line)
            set_run_font(run, 10.5, bold=(i == 0))


def build_document():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # --- Title ---
    add_heading(doc, "學界無人機救援挑戰賽 2026（ADRC 2026）", level=1)
    add_paragraph(
        doc,
        "模擬器編程任務 — 比賽當日參賽說明",
        size=12,
        bold=True,
        space_after=2,
    )
    add_paragraph(
        doc,
        "今日在 3D 模擬器中完成任務一與任務二，以積木程式控制無人機自主完成救災任務。",
        size=10.5,
        space_after=6,
    )

    # --- Flow ---
    add_heading(doc, "當日流程", level=2)
    add_table(
        doc,
        ["步驟", "內容"],
        [
            ("1", "開啟模擬器，選擇「競賽任務」"),
            ("2", "輸入比賽當日公佈的密碼解鎖正式任務"),
            ("3", "完成任務一 → 匯出程式 → 提交 Paperform"),
            ("4", "完成任務二 → 匯出程式 → 提交 Paperform"),
        ],
        col_widths_cm=[1.2, 14.8],
    )

    # --- Mission 1 ---
    add_heading(doc, "任務一：坍塌廢墟搜救", level=2)
    add_paragraph(
        doc,
        "目標：從起點（藍色箭嘴）沿可通行路網飛行，可選完成 3 處巡檢回報，最後在終點（綠色區域）降落完成情報交付。地圖 12×12 格，每格 150 cm。",
        size=10.5,
        space_after=3,
    )
    add_table(
        doc,
        ["項目", "分數"],
        [
            ("每處巡檢回報（可選，懸停約 3 秒）", "+100"),
            ("抵達終點並降落", "+200"),
            ("時間獎 ≤60s / ≤90s / ≤2min / ≤3min", "+500 / +350 / +200 / +80"),
        ],
        col_widths_cm=[11, 5],
    )
    add_paragraph(doc, "重要提示：", size=10.5, bold=True, space_after=2)
    add_bullet(doc, "須沿路網飛行，不可翻越建築。")
    add_bullet(doc, "禁用「飛至座標」積木（任務一專用限制）。")
    add_bullet(doc, "終點須用降落積木結算，懸停不能完成任務。")

    # --- Mission 2 ---
    add_heading(doc, "任務二：山火智能應對", level=2)
    add_paragraph(
        doc,
        "目標：起飛後反覆取水 → 滅火，依優先序撲滅全部 4 處火點。地圖 14×14 格，每格 150 cm。",
        size=10.5,
        space_after=3,
    )
    add_table(
        doc,
        ["火點（優先序）", "座標", "分數"],
        [
            ("A（最優先）", "(2, 12)", "+200"),
            ("B", "(4, 10)", "+150"),
            ("C", "(11, 5)", "+125"),
            ("D", "(12, 10)", "+100"),
            ("撲滅全部 4 處", "—", "額外 +200"),
        ],
        col_widths_cm=[4, 3.5, 3.5],
    )
    add_table(
        doc,
        ["資源規則", "說明"],
        [
            ("水量", "水箱每次僅 1 次水量；須到藍色水源用「取水」積木補給"),
            ("電量", "起飛 20 行；每次移動（go … cm）消耗 1 行"),
            ("充電站（可選）", "懸停 ≥3 秒，每站 +15 行（每站限一次）"),
            ("不計行", "轉向／取水／滅火／起降不消耗電量"),
        ],
        col_widths_cm=[3.5, 12.5],
    )
    add_paragraph(
        doc,
        "時間獎：≤4min +450｜≤6min +300｜≤8min +150｜≤10min +50",
        size=10.5,
        space_after=3,
    )
    add_paragraph(
        doc,
        "專用積木：COLLECT WATER 取水、RELEASE WATER 噴水滅火",
        size=10.5,
        space_after=4,
    )

    # --- General notes ---
    add_heading(doc, "共通注意事項", level=2)
    add_bullet(doc, "按「執行」後程式自動運行，比賽中不可手動操控無人機。")
    add_bullet(doc, "程式須具備自主導航能力，應善用距離感測器等積木，避免硬編碼固定路線。")
    add_bullet(doc, "可隨時按頂部「任務說明」重溫規則與計分。")
    add_bullet(doc, "完成後顯示成績與等級（任務一一等 ≥920；任務二一等 ≥950）。")
    add_bullet(doc, "試用關卡（免密碼）僅供練習，正式計分以競賽任務為準。")

    # --- Export & submit ---
    add_heading(doc, "程式匯出與成績提交", level=2)
    add_highlight_box(
        doc,
        [
            ("請務必完成以下步驟並提交程式！", True),
            "1. 完成任務並確認成績後，按頂部「顯示積木區」開啟程式面板",
            "2. 在積木工具列點選「匯出」（↓ 圖示）",
            "3. 瀏覽器下載 XML 檔：任務一 drone-blockly-mission-1.xml；任務二 drone-blockly-mission-2.xml",
            "4. 前往成績提交表格上傳：https://rfowp4av.paperform.co",
            "5. 兩個任務須分別匯出、分別提交；填寫隊伍資料並上傳對應 XML 檔",
        ],
    )

    # --- Footer ---
    add_paragraph(
        doc,
        "祝比賽順利！如有技術問題請舉手向工作人員查詢。",
        size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=0,
        color=RGBColor(0x55, 0x55, 0x55),
    )

    return doc


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
